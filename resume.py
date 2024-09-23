import streamlit as st
import json
import io
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER

def load_data():
    if 'resume_data' not in st.session_state:
        st.session_state.resume_data = {
            "personal_info": {},
            "summary": "",
            "professional_experience": [],
            "education": [],
            "awards": [],
            "core_competencies": {},
            "certifications": [],
            "publications": []
        }
    return st.session_state.resume_data

def save_data(data):
    st.session_state.resume_data = data

def edit_personal_info(data):
    st.subheader("Personal Information")
    data["personal_info"]["name"] = st.text_input("Name", data["personal_info"].get("name", ""))
    data["personal_info"]["email"] = st.text_input("Email", data["personal_info"].get("email", ""))
    data["personal_info"]["phone"] = st.text_input("Phone", data["personal_info"].get("phone", ""))
    data["personal_info"]["location"] = st.text_input("Location", data["personal_info"].get("location", ""))
    data["personal_info"]["linkedin"] = st.text_input("LinkedIn", data["personal_info"].get("linkedin", ""))
    data["personal_info"]["website"] = st.text_input("Website", data["personal_info"].get("website", ""))

def edit_summary(data):
    st.subheader("Summary")
    data["summary"] = st.text_area("Enter summary", data.get("summary", ""))

def edit_professional_experience(data):
    st.subheader("Professional Experience")
    for i, job in enumerate(data["professional_experience"]):
        with st.expander(f"{job['title']} at {job['company']}"):
            job["title"] = st.text_input(f"Job Title", job["title"], key=f"job_title_{i}")
            job["company"] = st.text_input(f"Company", job["company"], key=f"job_company_{i}")
            job["location"] = st.text_input(f"Location", job["location"], key=f"job_location_{i}")
            job["start_date"] = st.text_input(f"Start Date", job["start_date"], key=f"job_start_date_{i}")
            job["end_date"] = st.text_input(f"End Date", job["end_date"], key=f"job_end_date_{i}")
            
            # Responsibilities section
            st.write("Responsibilities:")
            if "responsibilities" not in job:
                job["responsibilities"] = []
            
            for j, resp in enumerate(job["responsibilities"]):
                col1, col2, col3 = st.columns([3, 1, 0.5])
                with col1:
                    job["responsibilities"][j]["content"] = st.text_input(f"Responsibility {j+1}", resp.get("content", ""), key=f"job_resp_content_{i}_{j}")
                with col2:
                    job["responsibilities"][j]["internal_name"] = st.text_input(f"Internal Name (optional)", resp.get("internal_name", ""), key=f"job_resp_name_{i}_{j}")
                with col3:
                    if st.button("Delete", key=f"del_resp_{i}_{j}"):
                        job["responsibilities"].pop(j)
                        st.rerun()
            
            if st.button("Add Responsibility", key=f"add_resp_{i}"):
                job["responsibilities"].append({"content": "", "internal_name": ""})
                st.rerun()
    
    if st.button("Add New Job"):
        data["professional_experience"].append({
            "title": "",
            "company": "",
            "location": "",
            "start_date": "",
            "end_date": "",
            "responsibilities": []
        })

def edit_education(data):
    st.subheader("Education")
    for i, edu in enumerate(data["education"]):
        with st.expander(f"{edu['degree']} from {edu['institution']}"):
            edu["institution"] = st.text_input(f"Institution", edu["institution"], key=f"edu_inst_{i}")
            edu["location"] = st.text_input(f"Location", edu["location"], key=f"edu_loc_{i}")
            edu["degree"] = st.text_input(f"Degree", edu["degree"], key=f"edu_degree_{i}")
    
    if st.button("Add New Education"):
        data["education"].append({
            "institution": "",
            "location": "",
            "degree": ""
        })

def edit_awards(data):
    st.subheader("Awards")
    for i, award in enumerate(data["awards"]):
        with st.expander(f"{award['name']}"):
            award["name"] = st.text_input(f"Award Name", award["name"], key=f"award_name_{i}")
            award["date"] = st.text_input(f"Date", award["date"], key=f"award_date_{i}")
    
    if st.button("Add New Award"):
        data["awards"].append({
            "name": "",
            "date": ""
        })

def edit_core_competencies(data):
    st.subheader("Core Competencies")
    for category in data["core_competencies"]:
        data["core_competencies"][category] = st.text_input(f"{category}", ", ".join(data["core_competencies"][category])).split(", ")
    
    new_category = st.text_input("Add new category")
    if new_category:
        data["core_competencies"][new_category] = []

def edit_certifications(data):
    st.subheader("Certifications")
    data["certifications"] = st.text_area("Certifications (one per line)", "\n".join(data["certifications"])).split("\n")

def edit_publications(data):
    st.subheader("Publications")
    for i, pub in enumerate(data["publications"]):
        with st.expander(f"{pub['title']}"):
            pub["title"] = st.text_input(f"Publication Title", pub["title"], key=f"pub_title_{i}")
            pub["publisher"] = st.text_input(f"Publisher", pub["publisher"], key=f"pub_publisher_{i}")
            pub["date"] = st.text_input(f"Date", pub["date"], key=f"pub_date_{i}")
    
    if st.button("Add New Publication"):
        data["publications"].append({
            "title": "",
            "publisher": "",
            "date": ""
        })

def display_resume(data):
    st.header("Resume Preview")
    st.subheader(data['personal_info'].get('name', ''))
    st.write(f"Email: {data['personal_info'].get('email', '')}")
    st.write(f"Phone: {data['personal_info'].get('phone', '')}")
    st.write(f"Location: {data['personal_info'].get('location', '')}")
    st.write(f"LinkedIn: {data['personal_info'].get('linkedin', '')}")
    st.write(f"Website: {data['personal_info'].get('website', '')}")
    
    st.subheader("Summary")
    st.write(data["summary"])
    
    st.subheader("Professional Experience")
    for job in data["professional_experience"]:
        st.write(f"**{job['title']} at {job['company']}**")
        st.write(f"{job['location']} | {job['start_date']} - {job['end_date']}")
        for resp in job["responsibilities"]:
            st.write(f"- {resp['content']}")
    
    st.subheader("Education")
    for edu in data["education"]:
        st.write(f"{edu['degree']} - {edu['institution']}, {edu['location']}")
    
    st.subheader("Awards")
    for award in data["awards"]:
        st.write(f"{award['name']} - {award['date']}")
    
    st.subheader("Core Competencies")
    for category, skills in data["core_competencies"].items():
        st.write(f"**{category}:** {', '.join(skills)}")
    
    st.subheader("Certifications")
    for cert in data["certifications"]:
        st.write(cert)
    
    st.subheader("Publications")
    for pub in data["publications"]:
        st.write(f"{pub['title']} - {pub['publisher']}, {pub['date']}")

def export_to_doc(data):
    doc = Document()
    
    # Personal Information
    doc.add_heading(data['personal_info'].get('name', ''), 0)
    doc.add_paragraph(f"Email: {data['personal_info'].get('email', '')}")
    doc.add_paragraph(f"Phone: {data['personal_info'].get('phone', '')}")
    doc.add_paragraph(f"Location: {data['personal_info'].get('location', '')}")
    doc.add_paragraph(f"LinkedIn: {data['personal_info'].get('linkedin', '')}")
    doc.add_paragraph(f"Website: {data['personal_info'].get('website', '')}")
    
    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(data["summary"])
    
    # Professional Experience
    doc.add_heading('Professional Experience', level=1)
    for job in data["professional_experience"]:
        doc.add_paragraph(f"{job['title']} at {job['company']}", style='Heading 2')
        doc.add_paragraph(f"{job['location']} | {job['start_date']} - {job['end_date']}")
        for resp in job["responsibilities"]:
            doc.add_paragraph(resp['content'], style='List Bullet')
    
    # Education
    doc.add_heading('Education', level=1)
    for edu in data["education"]:
        doc.add_paragraph(f"{edu['degree']} - {edu['institution']}, {edu['location']}")
    
    # Awards
    doc.add_heading('Awards', level=1)
    for award in data["awards"]:
        doc.add_paragraph(f"{award['name']} - {award['date']}")
    
    # Core Competencies
    doc.add_heading('Core Competencies', level=1)
    for category, skills in data["core_competencies"].items():
        doc.add_paragraph(f"{category}: {', '.join(skills)}")
    
    # Certifications
    doc.add_heading('Certifications', level=1)
    for cert in data["certifications"]:
        doc.add_paragraph(cert)
    
    # Publications
    doc.add_heading('Publications', level=1)
    for pub in data["publications"]:
        doc.add_paragraph(f"{pub['title']} - {pub['publisher']}, {pub['date']}")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def export_to_pdf(data):
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=letter)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER))
    story = []

    # Personal Information
    story.append(Paragraph(data['personal_info'].get('name', ''), styles['Title']))
    story.append(Paragraph(f"Email: {data['personal_info'].get('email', '')}", styles['Normal']))
    story.append(Paragraph(f"Phone: {data['personal_info'].get('phone', '')}", styles['Normal']))
    story.append(Paragraph(f"Location: {data['personal_info'].get('location', '')}", styles['Normal']))
    story.append(Paragraph(f"LinkedIn: {data['personal_info'].get('linkedin', '')}", styles['Normal']))
    story.append(Paragraph(f"Website: {data['personal_info'].get('website', '')}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Summary
    story.append(Paragraph('Summary', styles['Heading1']))
    story.append(Paragraph(data["summary"], styles['Normal']))
    story.append(Spacer(1, 12))

    # Professional Experience
    story.append(Paragraph('Professional Experience', styles['Heading1']))
    for job in data["professional_experience"]:
        story.append(Paragraph(f"{job['title']} at {job['company']}", styles['Heading2']))
        story.append(Paragraph(f"{job['location']} | {job['start_date']} - {job['end_date']}", styles['Normal']))
        for resp in job["responsibilities"]:
            story.append(Paragraph(f"• {resp['content']}", styles['Normal']))
        story.append(Spacer(1, 6))

    # Education
    story.append(Paragraph('Education', styles['Heading1']))
    for edu in data["education"]:
        story.append(Paragraph(f"{edu['degree']} - {edu['institution']}, {edu['location']}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Awards
    story.append(Paragraph('Awards', styles['Heading1']))
    for award in data["awards"]:
        story.append(Paragraph(f"{award['name']} - {award['date']}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Core Competencies
    story.append(Paragraph('Core Competencies', styles['Heading1']))
    for category, skills in data["core_competencies"].items():
        story.append(Paragraph(f"{category}: {', '.join(skills)}", styles['Normal']))
    story.append(Spacer(1, 12))

    # Certifications
    story.append(Paragraph('Certifications', styles['Heading1']))
    for cert in data["certifications"]:
        story.append(Paragraph(cert, styles['Normal']))
    story.append(Spacer(1, 12))

    # Publications
    story.append(Paragraph('Publications', styles['Heading1']))
    for pub in data["publications"]:
        story.append(Paragraph(f"{pub['title']} - {pub['publisher']}, {pub['date']}", styles['Normal']))

    doc.build(story)
    bio.seek(0)
    return bio

def main():
    st.title("Resume Editor")

    # File upload
    uploaded_file = st.file_uploader("Upload a JSON resume file", type="json")
    if uploaded_file is not None:
        st.session_state.resume_data = json.load(uploaded_file)

    data = load_data()

    # Sidebar for navigation
    st.sidebar.title("Navigation")
    page = st.sidebar.radio("Go to", ["Personal Info", "Summary", "Professional Experience", "Education", "Awards", "Core Competencies", "Certifications", "Publications", "Preview"])

    if page == "Personal Info":
        edit_personal_info(data)
    elif page == "Summary":
        edit_summary(data)
    elif page == "Professional Experience":
        edit_professional_experience(data)
    elif page == "Education":
        edit_education(data)
    elif page == "Awards":
        edit_awards(data)
    elif page == "Core Competencies":
        edit_core_competencies(data)
    elif page == "Certifications":
        edit_certifications(data)
    elif page == "Publications":
        edit_publications(data)
    elif page == "Preview":
        display_resume(data)

    save_data(data)

    # Export options
    st.sidebar.title("Export Options")
    if st.sidebar.button("Export as JSON"):
